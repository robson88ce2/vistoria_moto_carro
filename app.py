import streamlit as st
from PIL import Image
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import io
from datetime import datetime
import os
import tempfile
import requests

from reportlab.lib.pagesizes import A4  # usar A4 ao invés de Letter
from reportlab.pdfgen import canvas
from reportlab.lib.units import inch, mm
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Image as PDFImage
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_CENTER, TA_JUSTIFY


# ---- Funções auxiliares ----

def set_cell_border(cell, border_size=1):
    """Adiciona bordas à célula da tabela no Word."""
    tc = cell._element.tcPr
    tcBorders = OxmlElement('w:tcBorders')
    for border in ['top', 'left', 'bottom', 'right']:
        edge = OxmlElement(f'w:{border}')
        edge.set(qn('w:val'), 'single')
        edge.set(qn('w:sz'), str(border_size * 4))  # *4 para converter unidades adequadas
        edge.set(qn('w:space'), '0')
        edge.set(qn('w:color'), '000000')
        tcBorders.append(edge)
    tc.append(tcBorders)


def add_header_to_word(doc):
    """Cria cabeçalho do documento Word com informações institucionais."""
    header = doc.sections[0].header
    header_para = header.paragraphs[0]
    header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    header_text = [
        "GOVERNO DO ESTADO DO CEARÁ",
        "SECRETARIA DE SEGURANÇA PÚBLICA E DEFESA SOCIAL",
        "DEPARTAMENTO DE POLÍCIA JUDICIÁRIA DO INTERIOR NORTE",
        "4ª SECCIONAL",
        "DELEGACIA DE POLÍCIA CIVIL DE ITAPIPOCA"
    ]
    
    run = header_para.add_run("\n".join(header_text))
    run.bold = True
    run.font.size = Pt(10)
    
    # Linha horizontal ou espaço após cabeçalho
    paragraph = doc.add_paragraph()
    run2 = paragraph.add_run()
    run2.add_break()


def gerar_word(tipo, marca_modelo, cor, placa_ost, placa_verd, motor, chassi, fotos, data_atual):
    doc = Document()
    
    # Margens
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    # Cabeçalho institucional
    add_header_to_word(doc)
    
    # Título
    title = doc.add_heading('LAUDO DE VISTORIA VEICULAR', level=1)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Texto introdutório com data automática
    texto_intro = (
        f"Eu, Robson Oliveira de Sousa, Matrícula 300124-7-X, habilitado conforme o Curso de "
        f"Vistoria Veicular e Inclusão/Exclusão de Gravame de Roubo/Furto de veículos, realizado "
        f"conforme a Portaria Nº117/2024 – DG/AESP/CE, anexo I, publicado no Diário Oficial do Estado "
        f"| série 3 | ano XVI nº052 | Fortaleza, em 15 de março de 2024, constato, após Vistoria "
        f"realizada no pátio da Delegacia de Polícia Civil de Itapipoca em {data_atual}, a presença "
        f"de padrão dos sinais identificadores do veículo de acordo com o usual do fabricante, quais "
        f"sejam o número do motor, etiquetas e número VIN (chassi).\n\n"
    )
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    para.add_run(texto_intro)
    
    # Seção VEÍCULO VISTORIADO
    veiculo_heading = doc.add_heading('VEÍCULO VISTORIADO', level=2)
    veiculo_heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    tabela = doc.add_table(rows=0, cols=2)
    tabela.style = 'Table Grid'
    
    campos = [
        ("Tipo", tipo),
        ("Marca/Modelo", marca_modelo),
        ("Cor", cor),
        ("Placa Ostentada", placa_ost),
        ("Placa Verdadeira", placa_verd),
        ("Número do Motor", motor),
        ("Número do Chassi", chassi)
    ]
    for campo, valor in campos:
        row = tabela.add_row().cells
        # lado esquerdo negrito
        cell_left = row[0]
        cell_right = row[1]
        cell_left.text = campo
        # colocar negrito
        run_left = cell_left.paragraphs[0].runs[0]
        run_left.bold = True
        cell_right.text = valor or ""
        # Add bordas
        set_cell_border(cell_left)
        set_cell_border(cell_right)
    
    # Centralizar células
    for row in tabela.rows:
        for cell in row.cells:
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    
    # Imagens
    doc.add_heading('IMAGENS DOS ITENS VERIFICADOS', level=2).alignment = WD_ALIGN_PARAGRAPH.CENTER
    for legenda, img in fotos:
        doc.add_paragraph(legenda).alignment = WD_ALIGN_PARAGRAPH.CENTER
        img_io = io.BytesIO()
        img.save(img_io, format='PNG')
        img_io.seek(0)
        doc.add_picture(img_io, width=Inches(5))
        last_par = doc.paragraphs[-1]
        last_par.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()  # espaço
    
    # Conclusão, data e assinatura
    doc.add_paragraph("Sem mais para tratar, é o laudo de vistoria que segue.")
    conclusion_para = doc.add_paragraph()
    conclusion_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    conclusion_para.add_run(f"Itapipoca, {data_atual}")
    
    signature_para = doc.add_paragraph()
    signature_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    signature_para.add_run("\n\n\n_________________________________\nROBSON OLIVEIRA DE SOUSA\n OIP - Mat. 300124-7-X")
    
    return doc


# ---- Geração PDF com correções ----

def gerar_pdf(tipo, marca_modelo, cor, placa_ost, placa_verd, motor, chassi, fotos, data_atual):
    # criar temp
    temp_file = tempfile.NamedTemporaryFile(suffix='.pdf', delete=False)
    file_path = temp_file.name
    temp_file.close()
    
    doc = SimpleDocTemplate(
        file_path,
        pagesize=A4,
        rightMargin=72,
        leftMargin=72,
        topMargin=72,
        bottomMargin=72
    )
    
    styles = getSampleStyleSheet()
    # ajustar estilo já existente ou criar com nome diferente
    # não adicionar 'Title' novo se já existe
    if 'TitleCustom' not in styles.byName:
        styles.add(ParagraphStyle(name='TitleCustom', alignment=TA_CENTER, fontSize=16, leading=20))
    if 'Center' not in styles.byName:
        styles.add(ParagraphStyle(name='Center', alignment=TA_CENTER, fontSize=12))
    if 'Justify' not in styles.byName:
        styles.add(ParagraphStyle(name='Justify', alignment=TA_JUSTIFY, fontSize=12))
    if 'Header' not in styles.byName:
        styles.add(ParagraphStyle(name='Header', alignment=TA_CENTER, fontSize=10))
    
    elements = []
    
    # Cabeçalho seção institucional
    header_lines = [
        "GOVERNO DO ESTADO DO CEARÁ",
        "SECRETARIA DE SEGURANÇA PÚBLICA E DEFESA SOCIAL",
        "DEPARTAMENTO DE POLÍCIA JUDICIÁRIA DO INTERIOR NORTE",
        "4ª SECCIONAL",
        "DELEGACIA DE POLÍCIA CIVIL DE ITAPIPOCA"
    ]
    for line in header_lines:
        elements.append(Paragraph(line, styles['Header']))
    elements.append(Spacer(1, 0.2 * inch))
    
    elements.append(Paragraph("LAUDO DE VISTORIA VEICULAR", styles['TitleCustom']))
    elements.append(Spacer(1, 0.2 * inch))
    
    # Introdução
    intro = (
        f"Eu, Robson Oliveira de Sousa, Matrícula 300124-7-X, habilitado conforme o Curso de "
        f"Vistoria Veicular e Inclusão/Exclusão de Gravame de Roubo/Furto de veículos, realizado "
        f"conforme a Portaria Nº117/2024 – DG/AESP/CE, anexo I, publicado no Diário Oficial do Estado "
        f"| série 3 | ano XVI nº052 | Fortaleza, em 15 de março de 2024, constato, após Vistoria "
        f"realizada no pátio da Delegacia de Polícia Civil de Itapipoca em {data_atual}, a presença "
        f"de padrão dos sinais identificadores do veículo de acordo com o usual do fabricante, quais "
        f"sejam o número do motor, etiquetas e número VIN (chassi)."
    )
    elements.append(Paragraph(intro, styles['Justify']))
    elements.append(Spacer(1, 0.1 * inch))
    
    complementary = (
        "Contudo, após consulta no sistema informatizado, foi constatado que as placas ostentadas "
        "não condizem com os dados verdadeiros do veículo. Além disso, os QR Codes não geram resultado "
        "quando consultados no aplicativo VIO, indicando adulteração da referida placa."
    )
    elements.append(Paragraph(complementary, styles['Justify']))
    elements.append(Spacer(1, 0.2 * inch))
    
    elements.append(Paragraph("VEÍCULO VISTORIADO", styles['TitleCustom']))
    elements.append(Spacer(1, 0.1 * inch))
    
    # Dados do veículo
    dados = [
        f"<b>Tipo:</b> {tipo}",
        f"<b>Marca/Modelo:</b> {marca_modelo}",
        f"<b>Cor:</b> {cor}",
        f"<b>Placa Ostentada:</b> {placa_ost}",
        f"<b>Placa Verdadeira:</b> {placa_verd}",
        f"<b>Número do Motor:</b> {motor}",
        f"<b>Número do Chassi:</b> {chassi}"
    ]
    for d in dados:
        elements.append(Paragraph(d, styles['Center']))
        elements.append(Spacer(1, 0.05 * inch))
    elements.append(Spacer(1, 0.2 * inch))
    
    elements.append(Paragraph("IMAGENS DOS ITENS VERIFICADOS", styles['TitleCustom']))
    elements.append(Spacer(1, 0.1 * inch))
    
    # Inserir imagens
    for legenda, img in fotos:
        # imagem em memória
        img_bytes = io.BytesIO()
        img.save(img_bytes, format='PNG')
        img_bytes.seek(0)
        
        elements.append(Paragraph(f"<b>{legenda}</b>", styles['Center']))
        elements.append(Spacer(1, 0.05 * inch))
        # Mantendo proporção
        # largura fixa
        img_width = 4 * inch
        img_height = (img.height / img.width) * img_width
        elements.append(PDFImage(img_bytes, width=img_width, height=img_height))
        elements.append(Spacer(1, 0.2 * inch))
    
    # Conclusão
    elements.append(Paragraph("Sem mais para tratar, é o laudo de vistoria que segue.", styles['Justify']))
    elements.append(Spacer(1, 0.2 * inch))
    
    elements.append(Paragraph(f"Itapipoca, {data_atual}", styles['Center']))
    elements.append(Spacer(1, 0.5 * inch))
    elements.append(Paragraph("_________________________________", styles['Center']))
    elements.append(Paragraph("ROBSON OLIVEIRA DE SOUSA", styles['Center']))
    elements.append(Paragraph("OIP - Mat. 300124-7-X", styles['Center']))
    
    # Build PDF
    doc.build(elements)
    
    return file_path


# ---- Interface Streamlit ----

def main():
    st.set_page_config(page_title="Sistema de Vistoria Veicular", page_icon="🚗", layout="wide")
    st.title("📋 Sistema de Vistoria Veicular")
    
    # Dados do veículo
    st.subheader("Dados do Veículo")
    col1, col2 = st.columns(2)
    with col1:
        tipo = st.selectbox("Tipo do veículo", ["Motocicleta/Motoneta", "Automóvel"])
        marca_modelo = st.text_input("Marca/Modelo")
        cor = st.text_input("Cor")
        placa_ost = st.text_input("Placa Ostentada")
        placa_verd = st.text_input("Placa Verdadeira")
        motor = st.text_input("Número do Motor")
        chassi = st.text_input("Número do Chassi")
       
    with col2:
       
        
        st.subheader("Fotos")
        tipo_fotos = ["Foto do Motor", "Foto do Chassi", "Foto Consulta VIO"]
        if "Moto" in tipo:
            tipo_fotos += ["Foto do Veículo"]
        else:
            tipo_fotos += ["Foto da Traseira", "Foto da Dianteira", "Fotos das Etiquetas"]
        
        fotos = []
        for legenda in tipo_fotos:
            foto = st.camera_input(f"Tirar ou enviar {legenda}", key=legenda)
            if foto:
                fotos.append((legenda, Image.open(foto)))
    
    # Botões
    st.subheader("Gerar Documentos")
    col_word, col_pdf, col_fipe = st.columns(3)
    
    
    # Gerar Word
    if col_word.button("📄 Gerar Word"):
        if not marca_modelo or not placa_ost or not chassi:
            st.error("Por favor, preencha Marca/Modelo, Placa Ostentada e Chassi.")
        else:
            data_atual = datetime.now().strftime('%d/%m/%Y')
            with st.spinner("Gerando documento Word..."):
                doc = gerar_word(tipo, marca_modelo, cor, placa_ost, placa_verd, motor, chassi, fotos, data_atual)
                buffer = io.BytesIO()
                doc.save(buffer)
                buffer.seek(0)
                st.success("Documento Word gerado com sucesso!")
                st.download_button(
                    "📥 Baixar Word",
                    buffer,
                    file_name=f"vistoria_{placa_ost}_{datetime.now().strftime('%d%m%Y')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
    
    # Gerar PDF
    if col_pdf.button("📊 Gerar PDF"):
        if not marca_modelo or not placa_ost or not chassi:
            st.error("Por favor, preencha Marca/Modelo, Placa Ostentada e Chassi.")
        else:
            data_atual = datetime.now().strftime('%d/%m/%Y')
            with st.spinner("Gerando documento PDF..."):
                pdf_path = gerar_pdf(tipo, marca_modelo, cor, placa_ost, placa_verd, motor, chassi, fotos, data_atual)
                # ler bytes
                with open(pdf_path, "rb") as f:
                    pdf_bytes = f.read()
                st.success("Documento PDF gerado com sucesso!")
                st.download_button(
                    "📥 Baixar PDF",
                    pdf_bytes,
                    file_name=f"vistoria_{placa_ost}_{datetime.now().strftime('%d%m%Y')}.pdf",
                    mime="application/pdf"
                )
                try:
                    os.remove(pdf_path)
                except Exception:
                    pass


if __name__ == "__main__":
    main()
